"""Resume text extraction for PDF / DOCX / TXT. Best-effort; never raises."""
import os


def _parse_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        return ""


def _parse_docx(path: str) -> str:
    try:
        import docx  # python-docx

        document = docx.Document(path)
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception:
        return ""


def _parse_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except OSError:
        return ""


def parse_resume_file(path: str) -> str:
    """Return extracted plain text, or "" if the file can't be parsed."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext in (".docx", ".doc"):
        return _parse_docx(path)
    if ext in (".txt", ".md", ".text"):
        return _parse_txt(path)
    # Unknown extension: try text as a last resort.
    return _parse_txt(path)
