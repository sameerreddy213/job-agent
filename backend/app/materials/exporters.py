"""Export a list of (heading, body) sections to TXT / DOCX / PDF."""


def write_txt(path: str, sections: list[tuple[str, str]]) -> None:
    parts: list[str] = []
    for heading, body in sections:
        parts.append(heading.upper())
        parts.append("=" * len(heading))
        parts.append(body)
        parts.append("")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(parts))


def write_docx(path: str, sections: list[tuple[str, str]]) -> None:
    import docx

    document = docx.Document()
    for heading, body in sections:
        document.add_heading(heading, level=1)
        for line in body.split("\n"):
            document.add_paragraph(line)
    document.save(path)


def write_pdf(path: str, sections: list[tuple[str, str]]) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from xml.sax.saxutils import escape

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(path, pagesize=A4)
    flow = []
    for heading, body in sections:
        flow.append(Paragraph(escape(heading), styles["Heading1"]))
        for line in body.split("\n"):
            flow.append(Paragraph(escape(line) or "&nbsp;", styles["BodyText"]))
        flow.append(Spacer(1, 12))
    doc.build(flow)
